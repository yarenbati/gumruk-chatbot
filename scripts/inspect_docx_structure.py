"""
Milestone M1 - one-off analysis script.

Inspects the structure of the first official source document in data/raw/
(5326 Sayılı Kabahatler Kanunu) to inform the design of the future
src/ingest.py and src/chunk.py implementations.

This script is analysis tooling only. It is intentionally kept out of
src/ so it never gets mistaken for the production ingestion/chunking
pipeline (see AGENTS.md). It may stay in the repo for reproducibility.

Usage:
    python scripts/inspect_docx_structure.py

Notes on the source file
-------------------------
The source in data/raw/ is now a normalized-derivative ".docx" file (a
valid OOXML/zip package), converted from the originally obtained legacy
binary ".doc" using Microsoft Word (see docs/source-analysis-5326.md for
the history). This script reads it directly with `python-docx` — no
external CLI tool (antiword, LibreOffice, etc.) is used or required.

Two structures live outside `python-docx`'s high-level `Document.paragraphs`
API and are read via `python-docx`'s own package/part model instead:
  - the amendment-history table -> `Document.tables`
  - the amendment footnotes -> `word/footnotes.xml`, a part python-docx does
    not parse into `.paragraphs`; we locate it through `document.part.package`
    and parse it with `lxml` (already a python-docx dependency), not a new
    external tool.
"""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

import docx
from lxml import etree

# Windows consoles often default stdout to a legacy codepage (e.g. cp1252),
# which cannot represent all Turkish characters and breaks this script's
# output. Force UTF-8 so findings print reliably regardless of console.
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

RAW_DIR = Path(__file__).resolve().parent.parent / "data" / "raw"
WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W = WORD_NS["w"]

ARTICLE_PATTERN = re.compile(r"^Madde\s+(\d+(?:/[A-Z])?)\s*[-–—]\s*(.*)$")
EK_MADDE_PATTERN = re.compile(r"^Ek Madde\s+(\d+)\s*[-–—]\s*(.*)$")
GECICI_MADDE_PATTERN = re.compile(r"^Geçici Madde\s+(\d+)\s*[-–—]\s*(.*)$")
NUMBERED_PARAGRAPH_PATTERN = re.compile(r"^\((\d+)\)\s")
SECTION_HEADING_PATTERN = re.compile(
    r"^(BİRİNCİ|İKİNCİ|ÜÇÜNCÜ|DÖRDÜNCÜ|BEŞİNCİ|ALTINCI|YEDİNCİ|SEKİZİNCİ)"
    r"\s+(KISIM|BÖLÜM)$"
)

# Baseline counts from the previous antiword-based analysis (see
# docs/source-analysis-5326.md, "M1" section) — used to detect drift after
# switching to the python-docx/.docx source.
BASELINE = {
    "normal_articles": 45,
    "lettered_articles": 4,
    "ek_madde": 1,
    "gecici_madde": 3,
    "total_units": 53,
    "numbered_paragraphs": 105,
    "section_headings": 6,
    "footnote_definitions": 13,
}


def find_source_file() -> Path:
    candidates = sorted(RAW_DIR.glob("*.docx"))
    if not candidates:
        raise FileNotFoundError(
            f"No .docx source file found in {RAW_DIR}. Expected the "
            "normalized 5326 Sayılı Kabahatler Kanunu .docx source."
        )
    if len(candidates) > 1:
        print(
            f"WARNING: multiple .docx files found in {RAW_DIR}; "
            f"using the first: {candidates[0].name}"
        )
    return candidates[0]


def open_document(path: Path) -> docx.document.Document:
    print("=" * 70)
    print("STEP 1: open with python-docx")
    print("=" * 70)
    document = docx.Document(str(path))
    print(f"OK: python-docx opened '{path.name}' directly (no fallback used).")
    print(f"  paragraphs: {len(document.paragraphs)}, tables: {len(document.tables)}")
    return document


def analyze_page_numbers(path: Path) -> None:
    print()
    print("=" * 70)
    print("STEP 2: page-number reliability")
    print("=" * 70)
    with zipfile.ZipFile(path) as z:
        app_xml = z.read("docProps/app.xml").decode("utf-8")
    match = re.search(r"<Pages>(\d+)</Pages>", app_xml)
    if match:
        print(f"docProps/app.xml reports an application-computed page count: {match.group(1)}")
    print(
        "-> This is a document-level, print-layout-dependent statistic (a "
        "summary property, like Word's word count), not a per-article or "
        "per-paragraph page number. DOCX is a reflowable format; "
        "python-docx exposes no per-paragraph page number at all.\n"
        "CONCLUSION: 'legislation_number + article_no' (with paragraph_no "
        "when useful) remains the primary citation anchor, not page."
    )


def paragraph_texts(document: docx.document.Document) -> list[str]:
    return [p.text for p in document.paragraphs]


def analyze_paragraphs(texts: list[str]) -> None:
    print()
    print("=" * 70)
    print("STEP 3: paragraph count")
    print("=" * 70)
    non_empty = [t for t in texts if t.strip()]
    print(f"Total python-docx paragraphs: {len(texts)}")
    print(f"Non-empty paragraphs: {len(non_empty)}")


def find_article_starts(texts: list[str]) -> list[tuple[int, str, str]]:
    """Returns (paragraph_index, kind, article_no) for every article-like
    unit, in document order. kind is one of 'normal', 'lettered', 'ek',
    'gecici'."""
    starts: list[tuple[int, str, str]] = []
    for i, raw in enumerate(texts):
        t = raw.strip()
        if not t:
            continue
        m = EK_MADDE_PATTERN.match(t)
        if m:
            starts.append((i, "ek", m.group(1)))
            continue
        m = GECICI_MADDE_PATTERN.match(t)
        if m:
            starts.append((i, "gecici", m.group(1)))
            continue
        m = ARTICLE_PATTERN.match(t)
        if m:
            kind = "lettered" if "/" in m.group(1) else "normal"
            starts.append((i, kind, m.group(1)))
    return starts


def analyze_articles(texts: list[str], starts: list[tuple[int, str, str]]) -> None:
    print()
    print("=" * 70)
    print("STEP 4: article pattern detection")
    print("=" * 70)

    normal = [s for s in starts if s[1] == "normal"]
    lettered = [s for s in starts if s[1] == "lettered"]
    ek = [s for s in starts if s[1] == "ek"]
    gecici = [s for s in starts if s[1] == "gecici"]
    numbered_paragraphs = [t for t in texts if NUMBERED_PARAGRAPH_PATTERN.match(t.strip())]
    section_headings = [t for t in texts if SECTION_HEADING_PATTERN.match(t.strip())]

    print(f"Normal 'Madde N-' articles matched:        {len(normal)}")
    print(f"Lettered 'Madde N/A-' style articles:      {len(lettered)}")
    print(f"'Ek Madde N-' articles matched:             {len(ek)}")
    print(f"'Geçici Madde N-' articles matched:         {len(gecici)}")
    print(f"Numbered sub-paragraphs '(N) ...' matched:  {len(numbered_paragraphs)}")
    print(f"KISIM/BÖLÜM section headings matched:       {len(section_headings)}")
    print(f"Total article-like units:                   {len(starts)}")

    if lettered:
        print(f"  Lettered article numbers: {', '.join(s[2] for s in lettered)}")
    if ek:
        print(f"  Ek Madde numbers: {', '.join(s[2] for s in ek)}")
    if gecici:
        print(f"  Geçici Madde numbers: {', '.join(s[2] for s in gecici)}")


def extract_footnote_definitions(document: docx.document.Document) -> list[tuple[str, str]]:
    """Reads word/footnotes.xml via python-docx's own package/part model
    (document.part.package) and parses it with lxml. Not exposed by
    Document.paragraphs; no external tool involved."""
    package = document.part.package
    parts = [p for p in package.parts if p.partname == "/word/footnotes.xml"]
    if not parts:
        return []
    root = etree.fromstring(parts[0].blob)
    definitions = []
    for fn in root.findall("w:footnote", WORD_NS):
        fn_type = fn.get(f"{{{W}}}type")
        if fn_type in ("separator", "continuationSeparator"):
            continue
        fn_id = fn.get(f"{{{W}}}id")
        text = "".join(t.text or "" for t in fn.findall(".//w:t", WORD_NS)).strip()
        definitions.append((fn_id, text))
    return definitions


def count_inline_footnote_references(document: docx.document.Document) -> int:
    count = 0
    for p in document.paragraphs:
        count += len(p._p.findall(".//w:footnoteReference", WORD_NS))
    return count


def analyze_footnotes(document: docx.document.Document) -> None:
    print()
    print("=" * 70)
    print("STEP 5: amendment footnotes")
    print("=" * 70)
    inline_refs = count_inline_footnote_references(document)
    definitions = extract_footnote_definitions(document)
    print(f"Inline <w:footnoteReference> elements in body: {inline_refs}")
    print(f"Footnote definitions in word/footnotes.xml:    {len(definitions)}")
    print(
        "-> In this .docx, amendment notes are real Word footnotes (a "
        "separate word/footnotes.xml part), NOT inline '[N]' bracket text "
        "in the paragraph flow as antiword had rendered them. "
        "python-docx's Document.paragraphs does not expose footnote text; "
        "it must be read from the footnotes part directly (done above)."
    )
    if definitions:
        first_id, first_text = definitions[0]
        print(f"  Sample footnote [{first_id}]: {first_text[:100]}...")


def analyze_table(document: docx.document.Document) -> None:
    print()
    print("=" * 70)
    print("STEP 6: amendment-history table")
    print("=" * 70)
    tables = document.tables
    print(f"Tables detected via Document.tables: {len(tables)}")
    if not tables:
        return
    table = tables[0]
    print(f"Table 0: {len(table.rows)} rows x {len(table.columns)} columns")
    print(f"  Header row: {[c.text.strip() for c in table.rows[0].cells]}")
    print(f"  Data rows (excluding header): {len(table.rows) - 1}")


def turkish_char_check(texts: list[str], extra_text: str = "") -> None:
    print()
    print("=" * 70)
    print("STEP 7: Turkish character preservation check")
    print("=" * 70)
    full_text = "\n".join(texts) + "\n" + extra_text
    sample_chars = "ığĞşŞöÖüÜçÇİı"
    missing = [c for c in sample_chars if c not in full_text]
    if missing:
        print(f"WARNING: characters not observed in extracted text: {missing}")
    else:
        print("All sampled Turkish-specific characters (ığĞşŞöÖüÜçÇİı) are present.")
    replacement_chars = full_text.count("�")
    print(f"Unicode replacement characters ('\\ufffd') found: {replacement_chars}")


def get_article_block(
    texts: list[str], starts: list[tuple[int, str, str]], kind: str, article_no: str
) -> str | None:
    for idx, (para_i, s_kind, s_no) in enumerate(starts):
        if s_kind == kind and s_no == article_no:
            start_i = para_i
            end_i = starts[idx + 1][0] if idx + 1 < len(starts) else len(texts)
            block = [t for t in texts[start_i:end_i] if t.strip()]
            return "\n".join(block)
    return None


def show_samples(texts: list[str], starts: list[tuple[int, str, str]]) -> None:
    print()
    print("=" * 70)
    print("STEP 8: representative samples (conversion-integrity check)")
    print("=" * 70)

    targets = [
        ("normal", "1", "Madde 1 (long article with lettered sub-items)"),
        ("normal", "2", "Madde 2 (normal single-paragraph article)"),
        ("normal", "11", "Madde 11 (multi-paragraph article)"),
        ("lettered", "42/A", "Madde 42/A (lettered article)"),
        ("ek", "1", "Ek Madde 1"),
        ("gecici", "1", "Geçici Madde 1"),
    ]
    for kind, no, label in targets:
        block = get_article_block(texts, starts, kind, no)
        print(f"--- {label} ---")
        print(block if block is not None else "NOT FOUND")
        print()


def print_comparison(counts: dict[str, int]) -> None:
    print()
    print("=" * 70)
    print("STEP 9: comparison against antiword-based baseline")
    print("=" * 70)
    header = f"{'metric':<24}{'old (antiword)':<16}{'new (python-docx)':<20}{'match?'}"
    print(header)
    print("-" * len(header))
    for key, old_value in BASELINE.items():
        new_value = counts.get(key)
        match = "SAME" if new_value == old_value else "DIFFERENT"
        print(f"{key:<24}{old_value:<16}{new_value if new_value is not None else '?':<20}{match}")
    print()
    print(
        "Note: 'inline footnote references' and 'table rows' are not included "
        "above as direct numeric baseline comparisons — their *representation* "
        "changed (real XML footnotes/table vs. antiword's flattened, "
        "line-wrapped text), even where underlying content is unchanged. "
        "See docs/source-analysis-5326.md for the reconciled explanation."
    )


def main() -> None:
    source_path = find_source_file()
    print(f"Source file: {source_path.relative_to(RAW_DIR.parent.parent)}")
    print(f"File size: {source_path.stat().st_size} bytes")
    print()

    document = open_document(source_path)
    analyze_page_numbers(source_path)

    texts = paragraph_texts(document)
    analyze_paragraphs(texts)

    starts = find_article_starts(texts)
    analyze_articles(texts, starts)

    analyze_footnotes(document)
    analyze_table(document)

    footnote_defs = extract_footnote_definitions(document)
    footnote_text_blob = "\n".join(text for _, text in footnote_defs)
    turkish_char_check(texts, extra_text=footnote_text_blob)

    show_samples(texts, starts)

    normal = [s for s in starts if s[1] == "normal"]
    lettered = [s for s in starts if s[1] == "lettered"]
    ek = [s for s in starts if s[1] == "ek"]
    gecici = [s for s in starts if s[1] == "gecici"]
    numbered_paragraphs = [t for t in texts if NUMBERED_PARAGRAPH_PATTERN.match(t.strip())]
    section_headings = [t for t in texts if SECTION_HEADING_PATTERN.match(t.strip())]
    counts = {
        "normal_articles": len(normal),
        "lettered_articles": len(lettered),
        "ek_madde": len(ek),
        "gecici_madde": len(gecici),
        "total_units": len(starts),
        "numbered_paragraphs": len(numbered_paragraphs),
        "section_headings": len(section_headings),
        "footnote_definitions": len(footnote_defs),
    }
    print_comparison(counts)

    print()
    print("=" * 70)
    print("Done. See docs/source-analysis-5326.md for the written analysis.")
    print("=" * 70)


if __name__ == "__main__":
    main()
