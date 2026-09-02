"""
Production ingestion for legal source .docx files.

Scope (Milestone M2): read a real .docx file with python-docx, preserve
document order, and emit non-empty paragraphs with conservative whitespace
normalization. This module does NOT detect articles/Kısım/Bölüm structure,
does NOT build chunks, and does NOT alter legal wording — see
docs/source-analysis-5326.md for the parsing strategy planned for a later
milestone (src/chunk.py), and AGENTS.md for why that stays out of scope here.

Footnote references (M3A follow-up): Word footnote-reference marks
(<w:footnoteReference w:id="N"/>) are NOT part of python-docx's high-level
`paragraph.text` — see docs/source-analysis-5326.md §4. Since footnote
extraction is a source-extraction concern, it belongs here rather than in
src/chunk.py: each ExtractedParagraph additionally carries the raw
footnote-reference IDs found in its own OOXML, read via python-docx's own
paragraph element (`paragraph._p`), not a separate zip/XML tool. This module
still does not resolve footnote *text* or build the amendment-history model
— only traceable reference IDs are preserved.

CLI usage:
    python -m src.ingest data/raw/5326-kabahatler-kanunu.docx
"""

from __future__ import annotations

import argparse
import json
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

import docx
from docx.document import Document as DocxDocument

# Windows consoles often default stdout to a legacy codepage (e.g. cp1252),
# which cannot represent all Turkish/Unicode paths and breaks CLI output.
# Only reconfigure when running as a script; importing this module for tests
# should not mutate the interpreter's stdout.
if __name__ == "__main__":  # pragma: no cover - exercised via CLI, not tests
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

PROJECT_ROOT = Path(__file__).resolve().parent.parent
PROCESSED_DIR = PROJECT_ROOT / "data" / "processed"
SOURCE_MANIFEST_PATH = PROJECT_ROOT / "data" / "source_manifest.json"

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


@dataclass(frozen=True)
class ExtractedParagraph:
    """One non-empty paragraph extracted from a source .docx, in document order."""

    index: int
    text: str
    style_name: str
    footnote_reference_ids: list[int] | None = None


def load_docx(path: Path) -> DocxDocument:
    """Open a .docx file with python-docx.

    Treats `path` as immutable input: this function only reads it. Raises
    FileNotFoundError if missing, or whatever python-docx raises if the file
    is not a valid OOXML package (no external-tool fallback is attempted).
    """
    if not path.exists():
        raise FileNotFoundError(f"Source file not found: {path}")
    return docx.Document(str(path))


def normalize_paragraph_text(text: str) -> str:
    """Conservatively normalize a single paragraph's text.

    Collapses any embedded line breaks (python-docx renders explicit
    in-paragraph line breaks as '\\n') into single spaces, and trims
    leading/trailing whitespace. Does not touch internal spacing/tabs,
    wording, punctuation, or Turkish characters, and never merges text
    across separate paragraphs — each ExtractedParagraph still corresponds
    to exactly one python-docx paragraph.
    """
    unified = text.replace("\r\n", "\n").replace("\r", "\n")
    single_line = " ".join(unified.split("\n"))
    return single_line.strip()


def _extract_footnote_reference_ids(paragraph: Any) -> list[int] | None:
    """Read Word footnote-reference IDs embedded in this paragraph's OOXML.

    `paragraph.text` (python-docx's public API) never surfaces
    <w:footnoteReference> marks — they carry no run text. We instead walk
    the paragraph's own underlying element (`paragraph._p`, the lxml/oxml
    node python-docx already builds for every paragraph) and look up
    footnoteReference children directly. This uses only python-docx's own
    object model; no separate zip/XML tool or footnotes.xml lookup is
    needed just to get the reference IDs (unlike resolving footnote *text*,
    which does require reading word/footnotes.xml separately - out of scope
    here, see module docstring).
    """
    refs = paragraph._p.findall(".//w:footnoteReference", WORD_NS)
    ids = [int(ref.get(f"{{{WORD_NS['w']}}}id")) for ref in refs if ref.get(f"{{{WORD_NS['w']}}}id") is not None]
    return ids or None


def extract_paragraphs(document: DocxDocument) -> list[ExtractedParagraph]:
    """Extract non-empty paragraphs from a python-docx Document, in document order.

    Only body-level paragraphs (Document.paragraphs) are read. Table cell
    text and footnote *text* are out of scope for this ingestion pass (see
    docs/source-analysis-5326.md §4-5): they live outside the paragraph
    flow and require separate handling once article-level parsing exists.
    Footnote *reference IDs* (not their text) attached to each paragraph
    are preserved via `footnote_reference_ids`.

    Note: a paragraph whose only content is a footnote reference (no other
    visible text) would still be filtered out as "blank" here, and its
    reference ID would be lost — this does not occur anywhere in the real
    5326 source (every footnote reference sits inside a paragraph that also
    has substantial visible text), so it is not handled as a special case.
    """
    extracted: list[ExtractedParagraph] = []
    for index, paragraph in enumerate(document.paragraphs):
        normalized = normalize_paragraph_text(paragraph.text)
        if not normalized:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        footnote_reference_ids = _extract_footnote_reference_ids(paragraph)
        extracted.append(
            ExtractedParagraph(
                index=index,
                text=normalized,
                style_name=style_name,
                footnote_reference_ids=footnote_reference_ids,
            )
        )
    return extracted


def _load_source_metadata(source_path: Path) -> dict[str, Any]:
    """Look up this source file's entry in data/source_manifest.json.

    Matches by resolved file path against each entry's `local_file`, so no
    legal metadata (document_id, title, etc.) is hard-coded here. Returns an
    empty dict if the manifest is missing or has no matching entry.
    """
    if not SOURCE_MANIFEST_PATH.exists():
        return {}
    entries = json.loads(SOURCE_MANIFEST_PATH.read_text(encoding="utf-8"))
    resolved_source = source_path.resolve()
    for entry in entries:
        local_file = entry.get("local_file")
        if local_file and (PROJECT_ROOT / local_file).resolve() == resolved_source:
            return entry
    return {}


def ingest_document(path: Path) -> dict[str, Any]:
    """Run the M2 ingestion pipeline for one .docx source file.

    Returns a JSON-serializable dict: document_id, source_file,
    paragraph_count, paragraphs. Deterministic — running this twice on the
    same input produces an identical result. Does not detect articles or
    build chunks.
    """
    document = load_docx(path)
    paragraphs = extract_paragraphs(document)
    metadata = _load_source_metadata(path)
    return {
        "document_id": metadata.get("document_id"),
        "source_file": metadata.get("local_file", path.as_posix()),
        "paragraph_count": len(paragraphs),
        "paragraphs": [asdict(p) for p in paragraphs],
    }


def write_processed_json(result: dict[str, Any], output_path: Path) -> Path:
    """Write an ingest_document() result to `output_path` as UTF-8 JSON."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    return output_path


def default_output_path(source_path: Path) -> Path:
    """data/processed/<source-stem>.paragraphs.json for a given source file."""
    return PROCESSED_DIR / f"{source_path.stem}.paragraphs.json"


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        description="Ingest a legal source .docx into an intermediate paragraphs JSON."
    )
    parser.add_argument("source", type=Path, help="Path to the source .docx file")
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Output JSON path (default: data/processed/<stem>.paragraphs.json)",
    )
    args = parser.parse_args(argv)

    result = ingest_document(args.source)
    output_path = args.output or default_output_path(args.source)
    write_processed_json(result, output_path)

    print(f"Input file: {args.source}")
    print(f"Extracted paragraphs: {result['paragraph_count']}")
    print(f"Output path: {output_path}")


if __name__ == "__main__":
    main()
