"""Tests for src/ingest.py (M2 production .docx ingestion).

The core unit tests below build a tiny synthetic .docx fixture on the fly
(via python-docx + tmp_path) so they run on any machine, in CI, without
requiring data/raw/ — which is intentionally gitignored (see AGENTS.md:
never commit real legal source documents). They must always run, never skip.

A separate, clearly-named integration test at the bottom of this file
exercises the real 5326 Kabahatler Kanunu source and SKIPs (does not fail)
when that file isn't present locally.
"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

import docx
import pytest
from lxml import etree

from src import ingest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
REAL_SOURCE_PATH = PROJECT_ROOT / "data" / "raw" / "5326-kabahatler-kanunu.docx"

TURKISH_SAMPLE_CHARS = "ığüşöçİĞÜŞÖÇ"

# (text, style_name_or_None) — one blank paragraph is included on purpose to
# exercise blank-paragraph filtering.
SAMPLE_PARAGRAPHS: list[tuple[str, str | None]] = [
    ("KABAHATLER KANUNU", "Heading 1"),
    ("BİRİNCİ BÖLÜM", None),
    ("Madde 1- (1) Türkçe karakter testi: ığüşöçİĞÜŞÖÇ.", None),
    ("", None),
    ("Madde 2- (1) İkinci örnek hüküm.", None),
]


def _build_sample_docx(path: Path) -> Path:
    """Programmatically build a tiny representative legal-style .docx fixture."""
    document = docx.Document()
    for text, style in SAMPLE_PARAGRAPHS:
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)
    document.save(str(path))
    return path


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    return _build_sample_docx(tmp_path / "sample.docx")


def _file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _inject_footnote_reference(paragraph, footnote_id: int) -> None:
    """Test-only helper: append a <w:footnoteReference w:id="N"/> run to a
    paragraph's OOXML via raw lxml. python-docx's public API has no way to
    add footnotes, so this mirrors the structure real Word footnotes use
    (see docs/source-analysis-5326.md §4) without needing a full
    word/footnotes.xml part — only the reference ID is under test here, not
    footnote text/resolution.
    """
    w_ns = ingest.WORD_NS["w"]
    run = etree.SubElement(paragraph._p, f"{{{w_ns}}}r")
    ref = etree.SubElement(run, f"{{{w_ns}}}footnoteReference")
    ref.set(f"{{{w_ns}}}id", str(footnote_id))


# --- Core unit tests (self-contained, no data/raw/ dependency) -------------


def test_docx_opens_successfully(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    assert document is not None
    assert len(document.paragraphs) == len(SAMPLE_PARAGRAPHS)


def test_extracted_paragraph_count_greater_than_zero(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    assert len(paragraphs) > 0


def test_non_empty_paragraph_count_matches_fixture(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    expected = sum(1 for text, _ in SAMPLE_PARAGRAPHS if text.strip())
    assert len(paragraphs) == expected


def test_paragraph_objects_have_required_fields(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    first = paragraphs[0]
    assert isinstance(first.index, int)
    assert isinstance(first.text, str) and first.text
    assert isinstance(first.style_name, str)


def test_style_name_is_captured(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    assert paragraphs[0].text == "KABAHATLER KANUNU"
    assert paragraphs[0].style_name == "Heading 1"
    assert paragraphs[1].style_name == "Normal"


def test_paragraphs_without_footnotes_have_null_footnote_ids(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    assert all(p.footnote_reference_ids is None for p in paragraphs)


def test_footnote_reference_ids_are_preserved(tmp_path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("Madde 1- (1) Footnotelu test hükmü.")
    _inject_footnote_reference(document.paragraphs[0], footnote_id=1)
    path = tmp_path / "footnote_sample.docx"
    document.save(str(path))

    loaded = ingest.load_docx(path)
    paragraphs = ingest.extract_paragraphs(loaded)

    assert len(paragraphs) == 1
    assert paragraphs[0].footnote_reference_ids == [1]
    # The reference mark itself must never leak into the visible text.
    assert "[1]" not in paragraphs[0].text


def test_multiple_footnote_references_on_one_paragraph_are_preserved(tmp_path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("Madde 1- (1) Birden fazla dipnotlu hüküm.")
    _inject_footnote_reference(document.paragraphs[0], footnote_id=2)
    _inject_footnote_reference(document.paragraphs[0], footnote_id=3)
    path = tmp_path / "multi_footnote_sample.docx"
    document.save(str(path))

    loaded = ingest.load_docx(path)
    paragraphs = ingest.extract_paragraphs(loaded)
    assert paragraphs[0].footnote_reference_ids == [2, 3]


def test_footnote_reference_ids_survive_json_roundtrip(tmp_path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("Madde 1- (1) Footnotelu test hükmü.")
    _inject_footnote_reference(document.paragraphs[0], footnote_id=7)
    docx_path = tmp_path / "footnote_sample.docx"
    document.save(str(docx_path))

    result = ingest.ingest_document(docx_path)
    output_path = ingest.write_processed_json(result, tmp_path / "out.paragraphs.json")
    loaded = json.loads(output_path.read_text(encoding="utf-8"))

    assert loaded["paragraphs"][0]["footnote_reference_ids"] == [7]


def test_turkish_characters_survive_extraction(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    full_text = "\n".join(p.text for p in paragraphs)
    missing = [c for c in TURKISH_SAMPLE_CHARS if c not in full_text]
    assert not missing, f"missing Turkish characters: {missing}"


def test_blank_paragraphs_are_not_emitted(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    assert all(p.text.strip() for p in paragraphs)


def test_paragraph_order_matches_source(sample_docx_path: Path) -> None:
    document = ingest.load_docx(sample_docx_path)
    paragraphs = ingest.extract_paragraphs(document)
    expected_order = [text for text, _ in SAMPLE_PARAGRAPHS if text.strip()]
    assert [p.text for p in paragraphs] == expected_order


def test_paragraph_ordering_is_deterministic(sample_docx_path: Path) -> None:
    result_a = ingest.ingest_document(sample_docx_path)
    result_b = ingest.ingest_document(sample_docx_path)
    assert result_a["paragraphs"] == result_b["paragraphs"]
    indices = [p["index"] for p in result_a["paragraphs"]]
    assert indices == sorted(indices)


def test_output_json_can_be_produced(sample_docx_path: Path, tmp_path: Path) -> None:
    result = ingest.ingest_document(sample_docx_path)
    output_path = ingest.write_processed_json(result, tmp_path / "out.paragraphs.json")

    assert output_path.exists()
    loaded = json.loads(output_path.read_text(encoding="utf-8"))
    expected_count = sum(1 for text, _ in SAMPLE_PARAGRAPHS if text.strip())
    assert loaded["paragraph_count"] == result["paragraph_count"] == expected_count
    assert {"document_id", "source_file", "paragraph_count", "paragraphs"} <= loaded.keys()


def test_source_file_is_not_modified(sample_docx_path: Path, tmp_path: Path) -> None:
    before_hash = _file_hash(sample_docx_path)
    before_mtime = sample_docx_path.stat().st_mtime_ns

    result = ingest.ingest_document(sample_docx_path)
    ingest.write_processed_json(result, tmp_path / "out.paragraphs.json")

    assert _file_hash(sample_docx_path) == before_hash
    assert sample_docx_path.stat().st_mtime_ns == before_mtime


# --- Optional integration test: real 5326 source ----------------------------
# Clearly separated from the core unit tests above. This is the ONLY test in
# this file allowed to depend on data/raw/; it SKIPs (never fails the suite)
# when the real source isn't present locally, e.g. in CI or a fresh clone.


@pytest.mark.skipif(
    not REAL_SOURCE_PATH.exists(),
    reason=f"integration test: real source not present locally: {REAL_SOURCE_PATH}",
)
def test_integration_real_kabahatler_kanunu_source() -> None:
    """Source-validation integration test against the actual 5326 Kabahatler
    Kanunu .docx. Optional/local-only — see docs/source-analysis-5326.md."""
    result = ingest.ingest_document(REAL_SOURCE_PATH)
    assert result["paragraph_count"] > 0
    assert result["document_id"] == "5326_kabahatler_kanunu"
    texts = [p["text"] for p in result["paragraphs"]]
    assert any(t.startswith("Madde 2-") for t in texts)
